import random
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from typing import List, Dict, Tuple
from rich.progress import track

# Constants for color definitions, matching CMY BK printer colors
COLORS: Dict[str, RGBColor] = {
    'black': RGBColor(0, 0, 0),
    'cyan': RGBColor(0, 255, 255),
    'magenta': RGBColor(255, 0, 255),
    'yellow': RGBColor(255, 255, 0)
}

# Text proportions for colors
PROPORTIONS: Dict[str, float] = {
    'black': 0.4,
    'cyan': 0.2,
    'magenta': 0.2,
    'yellow': 0.2
}

TOTAL_WORDS: int = 10000  # Total number of words to be generated in the document
FONT_NAME: str = 'SimSun'  # Standard Chinese font (SimSun is often used as a replacement for SongTi)
FONT_SIZE: int = 36  # Font size (in points)
LINE_SPACING: Pt = Pt(36)  # Line spacing (in points)
MARGINS: Cm = Cm(0.2)  # Margin size (in centimeters)


class ColorfulTextDocument:
    def __init__(self, total_words: int, proportions: Dict[str, float], colors: Dict[str, RGBColor]):
        """
        Initializes the ColorfulTextDocument instance.
        
        :param total_words: int - Total number of words to include in the document.
        :param proportions: Dict[str, float] - Dictionary containing proportions of each color.
        :param colors: Dict[str, RGBColor] - Dictionary containing RGBColor instances for each color.
        """
        self.total_words: int = total_words  # Total number of words to generate
        self.proportions: Dict[str, float] = proportions  # Proportions of each color
        self.colors: Dict[str, RGBColor] = colors  # Color definitions
        self.document: Document = Document()  # Word document instance

    def set_document_margins(self) -> None:
        """
        Sets the margins for the document.
        This method adjusts the top, bottom, left, and right margins of all sections in the document.
        """
        sections = self.document.sections
        for section in sections:
            section.top_margin = MARGINS
            section.bottom_margin = MARGINS
            section.left_margin = MARGINS
            section.right_margin = MARGINS

    def generate_texts(self) -> List[Tuple[str, str]]:
        """
        Generates random text with colors based on specified proportions.
        
        :return: List[Tuple[str, str]] - A list of tuples, each containing a color and a randomly generated Chinese character.
        """
        texts: List[Tuple[str, str]] = []
        # Generate text for each color based on its proportion
        for color, proportion in self.proportions.items():
            count: int = int(self.total_words * proportion)  # Number of words to generate for this color
            for _ in range(count):
                texts.append((color, self._generate_random_character()))
        random.shuffle(texts)  # Shuffle the list to mix colors
        return texts

    def _generate_random_character(self) -> str:
        """
        Generates a random Chinese character.
        
        :return: str - A randomly generated Chinese character.
        """
        return chr(random.randint(0x4E00, 0x9FFF))  # Unicode range for common Chinese characters

    def add_text_to_document(self, texts: List[Tuple[str, str]]) -> None:
        """
        Adds the generated texts to the document with appropriate formatting.
        
        :param texts: List[Tuple[str, str]] - A list of tuples containing a color and text to add to the document.
        This method creates a paragraph in the document and adds each word with the specified color and formatting.
        """
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.line_spacing = LINE_SPACING  # Set line spacing for the paragraph

        # Add each word to the paragraph with the appropriate color and formatting
        for color, word in track(texts, description="Adding text to document..."):
            run = paragraph.add_run(word)
            run.font.color.rgb = self.colors[color]  # Set the color of the text
            run.font.name = FONT_NAME  # Set the font name
            run.font.size = Pt(FONT_SIZE)  # Set the font size

    def save_document(self, filename: str) -> None:
        """
        Saves the generated document with a specific filename.
        
        :param filename: str - Name of the file to save.
        This method saves the current state of the document to a file with the given filename.
        """
        self.document.save(filename)


def main() -> None:
    """
    Main function to execute the document generation.
    This function creates an instance of ColorfulTextDocument, generates the content, and saves it to a .docx file.
    """
    # Create an instance of ColorfulTextDocument with the specified parameters
    doc_generator = ColorfulTextDocument(TOTAL_WORDS, PROPORTIONS, COLORS)
    doc_generator.set_document_margins()  # Set document margins
    texts = doc_generator.generate_texts()  # Generate random colored texts
    doc_generator.add_text_to_document(texts)  # Add texts to the document
    doc_generator.save_document("Color_Print_Maintenance.docx")  # Save the document


if __name__ == "__main__":
    main()