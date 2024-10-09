import random
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from typing import List
from rich.progress import track

# Constants for color definitions, matching CMY BK printer colors
COLORS = {
    'black': RGBColor(0, 0, 0),
    'cyan': RGBColor(0, 255, 255),
    'magenta': RGBColor(255, 0, 255),
    'yellow': RGBColor(255, 255, 0)
}

# Text proportions for colors
PROPORTIONS = {
    'black': 0.4,
    'cyan': 0.2,
    'magenta': 0.2,
    'yellow': 0.2
}

TOTAL_WORDS = 10000
FONT_NAME = 'SimSun'  # Standard Chinese font (SimSun is often used as a replacement for SongTi)
FONT_SIZE = 14  # Placeholder for 'xiaochu', needs to be adjusted as per exact requirements
LINE_SPACING = Pt(36)
MARGINS = Cm(0.2)


class ColorfulTextDocument:
    def __init__(self, total_words: int, proportions: dict, colors: dict):
        """
        Initializes the ColorfulTextDocument instance.
        :param total_words: Total number of words to include in the document.
        :param proportions: Dictionary containing proportions of each color.
        :param colors: Dictionary containing RGBColor instances for each color.
        """
        self.total_words = total_words
        self.proportions = proportions
        self.colors = colors
        self.document = Document()

    def set_document_margins(self):
        """Sets the margins for the document."""
        sections = self.document.sections
        for section in sections:
            section.top_margin = MARGINS
            section.bottom_margin = MARGINS
            section.left_margin = MARGINS
            section.right_margin = MARGINS

    def generate_texts(self) -> List[str]:
        """Generates random text for each color based on proportions."""
        texts = []
        for color, proportion in self.proportions.items():
            count = int(self.total_words * proportion)
            for _ in range(count):
                texts.append((color, self._generate_random_word()))
        random.shuffle(texts)
        return texts

    def _generate_random_word(self) -> str:
        """Generates a random word of 3-5 characters."""
        return ''.join(random.choices('ABCDEFGHIJKLMNOPQRSTUVWXYZ', k=random.randint(3, 5)))

    def add_text_to_document(self, texts: List[str]):
        """
        Adds the generated texts to the document with appropriate formatting.
        :param texts: List of tuples containing color and text.
        """
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE)
        paragraph.paragraph_format.line_spacing = LINE_SPACING

        for color, word in track(texts, description="Adding text to document..."):
            run = paragraph.add_run(word + ' ')
            run.font.color.rgb = self.colors[color]
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE)

    def save_document(self, filename: str):
        """
        Saves the generated document with a specific filename.
        :param filename: Name of the file to save.
        """
        self.document.save(filename)


def main():
    """Main function to execute the document generation."""
    doc_generator = ColorfulTextDocument(TOTAL_WORDS, PROPORTIONS, COLORS)
    doc_generator.set_document_margins()
    texts = doc_generator.generate_texts()
    doc_generator.add_text_to_document(texts)
    doc_generator.save_document("Color_Print_Maintenance.docx")


if __name__ == "__main__":
    main()